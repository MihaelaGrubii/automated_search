import docx
import urllib.request
from bs4 import BeautifulSoup
import os.path
from os import path
from docx import Document

counter = 0


def addToFile(name, content, nbr):
    topic = name
    name = (name.replace(' ', '')[:25])
    for char in name:
        if char in " ?.!/;:":
            name = name.replace(char, '')

    name = name + '.txt'
    global counter

    for el in content:
        counter += 1
        if path.exists(name):
            f = open(name, 'a+')
        else:
            f = open(name, 'w+')
            f.write('<<<This are the top ' + str(nbr) +
                    ' results for your search!>>>\n')
            f.write('Topic: ' + topic.upper() + '\n')
        f.write('nbr.' + str(counter) + ' result: ' + el + '\n')
        f.close()


def add_hyperlink(paragraph, url, text, color, underline):
    """
    A function that places a hyperlink within a paragraph object.

    :param paragraph: The paragraph we are adding the hyperlink to.
    :param url: A string containing the required url
    :param text: The text displayed for the url
    :return: The hyperlink object
    """

    # This gets access to the document.xml.rels file and gets a new relation id value
    part = paragraph.part
    r_id = part.relate_to(
        url, docx.opc.constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True)

    # Create the w:hyperlink tag and add needed values
    hyperlink = docx.oxml.shared.OxmlElement('w:hyperlink')
    hyperlink.set(docx.oxml.shared.qn('r:id'), r_id, )

    # Create a w:r element
    new_run = docx.oxml.shared.OxmlElement('w:r')

    # Create a new w:rPr element
    rPr = docx.oxml.shared.OxmlElement('w:rPr')

    # Add color if it is given
    if not color is None:
        c = docx.oxml.shared.OxmlElement('w:color')
        c.set(docx.oxml.shared.qn('w:val'), color)
        rPr.append(c)

    # Remove underlining if it is requested
    if not underline:
        u = docx.oxml.shared.OxmlElement('w:u')
        u.set(docx.oxml.shared.qn('w:val'), 'none')
        rPr.append(u)

    # Join all the xml elements together add add the required text to the w:r element
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)

    paragraph._p.append(hyperlink)

    return hyperlink


def addToDocx(name, content, nbr):
    topic = name
    name = (name.replace(' ', '')[:25])
    for char in name:
        if char in " ?.!/;:":
            name = name.replace(char, '')

    name = name + '.docx'

    document = Document()
    row = nbr + 1  # Number of rows you want
    col = 2  # Number of collumns you want
    table = document.add_table(rows=row, cols=col)
    # set your style, look at the help documentation for more help
    table.style = 'LightShading-Accent1'
    for y in range(row):
        for x in range(col):
            cell = table.cell(y, x)
            if y == 0:
                if x == 0:
                    continue
                else:
                    cell.text = '<<This are the top ' + \
                        str(nbr) + ' results for your search!>>>\n' + \
                        'Topic: ' + topic.upper()
            else:
                if x == 0:
                    cell.text = 'nbr.' + str(y)
                else:
                    try:
                        soup = BeautifulSoup(
                            urllib.request.urlopen(content[y - 1]), 'lxml')
                        subj = str(soup.title.string)
                    except:
                        subj = "This is a main page without title!"
                    p = cell.add_paragraph()
                    hyperlink = add_hyperlink(
                        p, content[y - 1], subj, '8822ff', True)
                    cell.content = hyperlink
    document.save(name)
